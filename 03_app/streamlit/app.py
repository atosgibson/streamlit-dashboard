import streamlit as st
import pandas as pd
import unicodedata
from pathlib import Path
import altair as alt

# Relatório Word
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime

MAP_STATUS = Path("04_docs/escopo/mapeamento_status_financeiro.csv")
MAP_SERV   = Path("04_docs/escopo/mapeamento_servicos_autofill.csv")
SLA_CAD    = Path("04_docs/escopo/cadastro_sla.csv")

LOGO_PATH = Path("04_docs/logos/logo3.jpeg")

C_ORANGE = "#FC4C02"
C_GREEN  = "#23382C"
C_GRAY   = "#737373"
C_WHITE  = "#FFFFFF"

def br_money(x: float) -> str:
    try:
        x = float(x)
    except Exception:
        x = 0.0
    return f"{x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def norm_key(x) -> str:
    if x is None:
        return ""
    s = str(x).strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = " ".join(s.split())
    return s

def load_mappings():
    ms = pd.read_csv(MAP_STATUS)
    ms["status_key"] = ms["STATUS"].apply(norm_key)

    mserv = pd.read_csv(MAP_SERV)
    serv_col = None
    for c in mserv.columns:
        if "SERV" in str(c).upper():
            serv_col = c
            break
    if serv_col is None:
        raise ValueError(f"Não achei coluna SERVIÇO no mapeamento. Colunas: {mserv.columns.tolist()}")

    mserv["servico_key"] = mserv[serv_col].apply(norm_key)

    sla = pd.read_csv(SLA_CAD)
    sla["k_cli"]  = sla["cliente"].apply(norm_key)
    sla["k_tipo"] = sla["tipo_servico"].apply(norm_key)
    sla["sla_dias"] = pd.to_numeric(sla["sla_dias"], errors="coerce")
    return ms, mserv, sla

def clean_excel(uploaded_file) -> pd.DataFrame:
    df = pd.read_excel(uploaded_file, header=1)
    df = df.loc[:, ~df.columns.astype(str).str.contains(r"^Unnamed")]
    df.columns = [str(c).strip() for c in df.columns]

    text_cols = ["CLIENTE", "PONTO", "UF", "GESTOR", "CHAMADO", "OS", "SERVIÇO", "STATUS"]
    for c in text_cols:
        if c in df.columns:
            df[c] = df[c].astype("string").str.strip()

    for c in ["AUTORIZAÇÃO", "TÉRMINO"]:
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors="coerce")

    df["receita"] = pd.to_numeric(df.get("RECEITA", 0), errors="coerce").fillna(0)
    df["sla_dias"] = (df["TÉRMINO"] - df["AUTORIZAÇÃO"]).dt.days
    df["mes_autorizacao"] = df["AUTORIZAÇÃO"].dt.to_period("M").astype(str)
    return df

def apply_enrichment(df: pd.DataFrame, ms: pd.DataFrame, mserv: pd.DataFrame, sla: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()

    df["status_key"]  = df["STATUS"].apply(norm_key)
    df["servico_key"] = df["SERVIÇO"].apply(norm_key)
    df["k_cli"]       = df["CLIENTE"].apply(norm_key)

    df = df.merge(ms[["status_key", "billing_status"]], on="status_key", how="left")
    df["billing_status"] = df["billing_status"].fillna("NAO_MAPEADO")

    df = df.merge(mserv[["servico_key", "tipo_servico", "categoria"]], on="servico_key", how="left")
    df["tipo_servico"] = df["tipo_servico"].fillna("NAO_MAPEADO")
    df["categoria"] = df["categoria"].fillna("NAO_MAPEADO")
    df["k_tipo"] = df["tipo_servico"].apply(norm_key)

    sla_especifico = sla[sla["cliente"] != "*"].copy()
    sla_padrao = sla[sla["cliente"] == "*"].copy()

    df = df.merge(
        sla_especifico[["k_cli", "k_tipo", "sla_dias"]].rename(columns={"sla_dias": "sla_meta_dias"}),
        on=["k_cli", "k_tipo"],
        how="left"
    )
    df = df.merge(
        sla_padrao[["k_tipo", "sla_dias"]].rename(columns={"sla_dias": "sla_meta_padrao"}),
        on="k_tipo",
        how="left"
    )
    df["sla_meta_dias"] = df["sla_meta_dias"].fillna(df["sla_meta_padrao"])
    df = df.drop(columns=["sla_meta_padrao"], errors="ignore")

    def sla_result(row):
        if pd.isna(row["sla_dias"]) or pd.isna(row["sla_meta_dias"]):
            return "SEM_DADO"
        return "DENTRO" if row["sla_dias"] <= row["sla_meta_dias"] else "FORA"

    df["sla_resultado"] = df.apply(sla_result, axis=1)
    return df.drop(columns=["status_key", "servico_key"], errors="ignore")

def gerar_relatorio_docx(df_filtrado: pd.DataFrame, filtros: dict) -> bytes:
    doc = Document()

    # Logo
    if LOGO_PATH.exists():
        try:
            doc.add_picture(str(LOGO_PATH), width=Inches(1.4))
        except Exception:
            pass

    doc.add_heading("Relatório Executivo • MJ Engenharia", level=1)
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # Filtros aplicados
    doc.add_heading("Escopo do relatório (filtros aplicados)", level=2)

    def fmt_lista(nome, lista):
        if not lista:
            return f"- {nome}: (nenhum selecionado)"
        if len(lista) <= 10:
            return f"- {nome}: " + ", ".join([str(x) for x in lista])
        return f"- {nome}: {len(lista)} selecionados (ex.: " + ", ".join([str(x) for x in lista[:5]]) + ", ...)"

    doc.add_paragraph(fmt_lista("Clientes", filtros.get("clientes", [])))
    doc.add_paragraph(fmt_lista("Tipos de serviço", filtros.get("tipos", [])))
    doc.add_paragraph(fmt_lista("Meses (Autorização)", filtros.get("meses", [])))

    # Faixa real de datas (autorização)
    if "AUTORIZAÇÃO" in df_filtrado.columns:
        dmin = pd.to_datetime(df_filtrado["AUTORIZAÇÃO"], errors="coerce").min()
        dmax = pd.to_datetime(df_filtrado["AUTORIZAÇÃO"], errors="coerce").max()
        if pd.notna(dmin) and pd.notna(dmax):
            doc.add_paragraph(f"- Faixa de Autorização no recorte: {dmin.strftime('%d/%m/%Y')} a {dmax.strftime('%d/%m/%Y')}")

    # KPIs
    total_reg = len(df_filtrado)
    chamados = int(df_filtrado["CHAMADO"].nunique()) if "CHAMADO" in df_filtrado.columns else 0
    receita_total = float(df_filtrado["receita"].sum()) if "receita" in df_filtrado.columns else 0.0
    pendencias = float(
        df_filtrado.loc[df_filtrado["billing_status"].isin(["PENDENTE_FATURAMENTO","FATURADO_PENDENTE"]), "receita"].sum()
    ) if ("billing_status" in df_filtrado.columns and "receita" in df_filtrado.columns) else 0.0
    fora_sla = int((df_filtrado.get("sla_resultado") == "FORA").sum()) if "sla_resultado" in df_filtrado.columns else 0

    # 1
    doc.add_heading("1) Qual o volume de chamados no recorte atual?", level=2)
    doc.add_paragraph(f"- Registros: {total_reg}")
    doc.add_paragraph(f"- Chamados únicos: {chamados}")
    doc.add_paragraph(f"- Clientes no recorte: {df_filtrado['CLIENTE'].nunique() if 'CLIENTE' in df_filtrado.columns else 0}")

    # 2
    doc.add_heading("2) Quais clientes mais acionaram a MJ (Top 10) e participação?", level=2)
    if "CLIENTE" in df_filtrado.columns and "CHAMADO" in df_filtrado.columns and total_reg > 0:
        top_cli = (df_filtrado.dropna(subset=["CLIENTE","CHAMADO"])
                   .groupby("CLIENTE")["CHAMADO"].nunique()
                   .reset_index(name="qtd_chamados")
                   .sort_values("qtd_chamados", ascending=False)
                   .head(10))
        total_ch = max(1, int(df_filtrado["CHAMADO"].nunique()))
        for _, r in top_cli.iterrows():
            pct = (int(r["qtd_chamados"]) / total_ch) * 100
            doc.add_paragraph(f"- {r['CLIENTE']}: {int(r['qtd_chamados'])} chamados ({pct:.1f}%)")
    else:
        doc.add_paragraph("Sem dados suficientes para cálculo.")

    # 3
    doc.add_heading("3) Quais tipos de serviço mais demandados (Top 10)?", level=2)
    if "tipo_servico" in df_filtrado.columns and "OS" in df_filtrado.columns:
        top_tipo = (df_filtrado.groupby("tipo_servico")["OS"].count()
                    .reset_index(name="qtd")
                    .sort_values("qtd", ascending=False)
                    .head(10))
        for _, r in top_tipo.iterrows():
            doc.add_paragraph(f"- {r['tipo_servico']}: {int(r['qtd'])} registros")
    else:
        doc.add_paragraph("Sem dados suficientes para cálculo.")

    # 4
    doc.add_heading("4) Como está o funil financeiro no recorte atual?", level=2)
    doc.add_paragraph(f"- Receita total: R$ {br_money(receita_total)}")
    if "billing_status" in df_filtrado.columns and "receita" in df_filtrado.columns:
        fin = (df_filtrado.groupby("billing_status")["receita"].sum()
               .reset_index(name="receita_total")
               .sort_values("receita_total", ascending=False))
        for _, r in fin.iterrows():
            doc.add_paragraph(f"- {r['billing_status']}: R$ {br_money(float(r['receita_total']))}")

    # 5
    doc.add_heading("5) Qual o valor em pendência e onde está concentrado?", level=2)
    doc.add_paragraph(f"- Pendências (A faturar + A receber): R$ {br_money(pendencias)}")
    if "billing_status" in df_filtrado.columns and "receita" in df_filtrado.columns and "CLIENTE" in df_filtrado.columns:
        pend = df_filtrado[df_filtrado["billing_status"].isin(["PENDENTE_FATURAMENTO","FATURADO_PENDENTE"])].copy()
        pend_cli = (pend.groupby("CLIENTE")["receita"].sum()
                    .reset_index(name="pendencia")
                    .sort_values("pendencia", ascending=False)
                    .head(5))
        if len(pend_cli) == 0:
            doc.add_paragraph("- Sem pendências no recorte atual.")
        else:
            doc.add_paragraph("Top 5 clientes por pendência:")
            for _, r in pend_cli.iterrows():
                doc.add_paragraph(f"- {r['CLIENTE']}: R$ {br_money(float(r['pendencia']))}")

    # 6
    doc.add_heading("6) Qual a performance de SLA (dentro/fora) no recorte atual?", level=2)
    if "sla_resultado" in df_filtrado.columns and total_reg > 0:
        dentro = int((df_filtrado["sla_resultado"] == "DENTRO").sum())
        fora = int((df_filtrado["sla_resultado"] == "FORA").sum())
        sem = int((df_filtrado["sla_resultado"] == "SEM_DADO").sum())
        doc.add_paragraph(f"- Dentro: {dentro} ({(dentro/total_reg*100):.1f}%)")
        doc.add_paragraph(f"- Fora: {fora} ({(fora/total_reg*100):.1f}%)")
        doc.add_paragraph(f"- Sem dado: {sem} ({(sem/total_reg*100):.1f}%)")
    else:
        doc.add_paragraph("Sem dados suficientes para cálculo de SLA.")

    # 7
    doc.add_heading("7) Quais tipos de serviço mais estouram SLA? (Top 5)", level=2)
    if "sla_resultado" in df_filtrado.columns and "tipo_servico" in df_filtrado.columns:
        sla_tipo = (df_filtrado.assign(is_fora=df_filtrado["sla_resultado"].eq("FORA"))
                    .groupby("tipo_servico")["is_fora"].sum()
                    .reset_index(name="fora")
                    .sort_values("fora", ascending=False)
                    .head(5))
        if len(sla_tipo) == 0 or int(sla_tipo["fora"].sum()) == 0:
            doc.add_paragraph("Sem casos fora do SLA no recorte atual.")
        else:
            for _, r in sla_tipo.iterrows():
                doc.add_paragraph(f"- {r['tipo_servico']}: {int(r['fora'])} fora")
    else:
        doc.add_paragraph("Sem dados suficientes para cálculo.")

    # 8
    doc.add_heading("8) Quais casos críticos (fora do SLA) exigem ação imediata? (Top 10)", level=2)
    if "sla_resultado" in df_filtrado.columns:
        crit = df_filtrado[df_filtrado["sla_resultado"] == "FORA"].copy()
        if len(crit) == 0:
            doc.add_paragraph("Nenhum caso fora do SLA no recorte atual.")
        else:
            crit["sla_dias_n"] = pd.to_numeric(crit.get("sla_dias", 0), errors="coerce")
            crit["sla_meta_n"] = pd.to_numeric(crit.get("sla_meta_dias", 0), errors="coerce")
            crit["atraso_dias"] = (crit["sla_dias_n"] - crit["sla_meta_n"]).fillna(0)
            crit = crit.sort_values("atraso_dias", ascending=False).head(10)

            for _, r in crit.iterrows():
                doc.add_paragraph(
                    f"- Cliente: {r.get('CLIENTE','')} | Ponto: {r.get('PONTO','')} | Tipo: {r.get('tipo_servico','')} | "
                    f"Atraso: {int(r.get('atraso_dias',0))} dias | OS: {r.get('OS','')} | Chamado: {r.get('CHAMADO','')}"
                )
    else:
        doc.add_paragraph("Sem dados suficientes para cálculo.")

    # 9
    doc.add_heading("9) Qual o tempo de ciclo médio/mediano por tipo de serviço?", level=2)
    if "tipo_servico" in df_filtrado.columns and "sla_dias" in df_filtrado.columns:
        ciclo = (df_filtrado.groupby("tipo_servico")["sla_dias"]
                 .agg(media="mean", mediana="median")
                 .reset_index()
                 .sort_values("media", ascending=False))
        for _, r in ciclo.head(10).iterrows():
            doc.add_paragraph(f"- {r['tipo_servico']}: média {r['media']:.1f} dias | mediana {r['mediana']:.1f} dias")
    else:
        doc.add_paragraph("Sem dados suficientes para cálculo.")

    # 10
    doc.add_heading("10) Resumo executivo: 3 ações recomendadas", level=2)
    doc.add_paragraph("- Ação 1: Priorizar faturamento dos itens em PENDENTE_FATURAMENTO.")
    doc.add_paragraph("- Ação 2: Cobrança/recebimento dos itens em FATURADO_PENDENTE.")
    doc.add_paragraph("- Ação 3: Revisar causas de FORA do SLA nos tipos com maior incidência.")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------------- UI ----------------
st.set_page_config(page_title="MJ Engenharia • Dashboard", layout="wide")

st.markdown(f"""
<style>
    .block-container {{ padding-top: 1.2rem; }}
    header {{ visibility: hidden; }}
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}

    .mj-card {{
        background: #0e0e0e;
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 18px;
        padding: 16px 16px;
    }}
    .mj-kpi-title {{
        color: {C_GRAY};
        font-size: 0.85rem;
        margin-bottom: 6px;
    }}
    .mj-kpi-value {{
        color: {C_WHITE};
        font-size: 1.6rem;
        font-weight: 700;
        line-height: 1.2;
    }}
    .mj-badge {{
        display: inline-block;
        padding: 6px 10px;
        border-radius: 999px;
        background: rgba(252, 76, 2, 0.12);
        border: 1px solid rgba(252, 76, 2, 0.35);
        color: {C_ORANGE};
        font-size: 0.85rem;
        font-weight: 600;
    }}
</style>
""", unsafe_allow_html=True)

h1, h2 = st.columns([1, 4])
with h1:
    if LOGO_PATH.exists():
        st.image(str(LOGO_PATH), width=110)
with h2:
    st.markdown("## MJ Engenharia")
    st.markdown('<span class="mj-badge">Soluções inteligentes • Dashboard</span>', unsafe_allow_html=True)

st.write("")
uploaded = st.file_uploader("Carregar planilha (.xlsx)", type=["xlsx"])
if uploaded is None:
    st.info("Anexe a planilha para atualizar os indicadores.")
    st.stop()

missing = [str(p) for p in [MAP_STATUS, MAP_SERV, SLA_CAD] if not p.exists()]
if missing:
    st.error("Faltam arquivos de configuração:\n- " + "\n- ".join(missing))
    st.stop()

ms, mserv, sla = load_mappings()
raw = clean_excel(uploaded)
df = apply_enrichment(raw, ms, mserv, sla)

st.sidebar.markdown("## Filtros")
clientes = sorted(df["CLIENTE"].dropna().unique().tolist())
tipos = sorted(df["tipo_servico"].dropna().unique().tolist())
meses = sorted(df["mes_autorizacao"].dropna().unique().tolist())

cli_sel = st.sidebar.multiselect("Cliente", clientes, default=clientes)
tipo_sel = st.sidebar.multiselect("Tipo de serviço", tipos, default=tipos)
mes_sel = st.sidebar.multiselect("Mês (Autorização)", meses, default=meses[-3:] if len(meses) >= 3 else meses)

f = df.copy()
if cli_sel:  f = f[f["CLIENTE"].isin(cli_sel)]
if tipo_sel: f = f[f["tipo_servico"].isin(tipo_sel)]
if mes_sel:  f = f[f["mes_autorizacao"].isin(mes_sel)]

total_linhas = len(f)
chamados_unicos = int(f["CHAMADO"].nunique())
receita_total = float(f["receita"].sum())
pendencias = float(f.loc[f["billing_status"].isin(["PENDENTE_FATURAMENTO", "FATURADO_PENDENTE"]), "receita"].sum())
fora_sla = int((f["sla_resultado"] == "FORA").sum())

k1, k2, k3, k4, k5 = st.columns(5)
for col, title, value in [
    (k1, "Registros", f"{total_linhas}"),
    (k2, "Chamados únicos", f"{chamados_unicos}"),
    (k3, "Receita total (R$)", br_money(receita_total)),
    (k4, "Pendências (R$)", br_money(pendencias)),
    (k5, "Fora do SLA", f"{fora_sla}"),
]:
    with col:
        st.markdown(f"""
        <div class="mj-card">
          <div class="mj-kpi-title">{title}</div>
          <div class="mj-kpi-value">{value}</div>
        </div>
        """, unsafe_allow_html=True)

st.write("")
tab1, tab2, tab3, tab4 = st.tabs(["📊 Visão geral", "💰 Financeiro", "⏱️ SLA", "📄 Base"])

with tab1:
    st.write("")
    if st.button("📄 Gerar relatório Word (recorte atual)"):
        filtros = {"clientes": cli_sel, "tipos": tipo_sel, "meses": mes_sel}
        doc_bytes = gerar_relatorio_docx(f, filtros)
        st.download_button(
            "⬇️ Baixar relatório (.docx)",
            data=doc_bytes,
            file_name="Relatorio_MJ_Dashboard.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    st.write("")

    cA, cB = st.columns([1.2, 1])

    top_clientes = (f.dropna(subset=["CLIENTE", "CHAMADO"])
                     .groupby("CLIENTE")["CHAMADO"].nunique()
                     .reset_index(name="qtd_chamados")
                     .sort_values("qtd_chamados", ascending=False)
                     .head(10))

    chart_cli = alt.Chart(top_clientes).mark_bar().encode(
        y=alt.Y("CLIENTE:N", sort="-x", title="Cliente"),
        x=alt.X("qtd_chamados:Q", title="Chamados"),
        color=alt.value(C_GREEN),
        tooltip=["CLIENTE", "qtd_chamados"]
    ).properties(height=320)

    dm = (f.groupby("mes_autorizacao")["OS"].count()
          .reset_index(name="qtd")
          .sort_values("mes_autorizacao"))

    chart_mes = alt.Chart(dm).mark_line(point=True).encode(
        x=alt.X("mes_autorizacao:N", title="Mês"),
        y=alt.Y("qtd:Q", title="Demandas"),
        color=alt.value(C_ORANGE),
        tooltip=["mes_autorizacao", "qtd"]
    ).properties(height=320)

    with cA:
        st.subheader("Chamados por cliente (Top 10)")
        st.altair_chart(chart_cli, use_container_width=True)
    with cB:
        st.subheader("Demandas por mês (Autorização)")
        st.altair_chart(chart_mes, use_container_width=True)

    st.subheader("Resumo por tipo de serviço")
    tipo_tbl = (f.groupby("tipo_servico")
                 .agg(qtd=("OS", "count"), receita=("receita", "sum"))
                 .reset_index()
                 .sort_values("qtd", ascending=False))
    st.dataframe(tipo_tbl, use_container_width=True, height=260)

with tab2:
    st.subheader("Receita por status")
    fin = (f.groupby("billing_status")["receita"].sum()
           .reset_index(name="receita_total")
           .sort_values("receita_total", ascending=False))

    chart_fin = alt.Chart(fin).mark_bar().encode(
        x=alt.X("billing_status:N", title="Status"),
        y=alt.Y("receita_total:Q", title="Receita (R$)"),
        color=alt.value(C_ORANGE),
        tooltip=["billing_status", alt.Tooltip("receita_total:Q", format=",.2f")]
    ).properties(height=320)

    st.altair_chart(chart_fin, use_container_width=True)
    st.dataframe(fin, use_container_width=True, height=220)

with tab3:
    st.subheader("SLA dentro/fora por tipo")
    sla_t = (f.groupby("tipo_servico")["sla_resultado"]
             .value_counts(dropna=False)
             .unstack(fill_value=0)
             .reset_index())

    for col in ["DENTRO", "FORA", "SEM_DADO"]:
        if col not in sla_t.columns:
            sla_t[col] = 0

    sla_long = sla_t.melt("tipo_servico", value_vars=["DENTRO", "FORA", "SEM_DADO"],
                          var_name="resultado", value_name="qtd")

    color_scale = alt.Scale(domain=["DENTRO","FORA","SEM_DADO"],
                            range=[C_GREEN, C_ORANGE, C_GRAY])

    chart_sla = alt.Chart(sla_long).mark_bar().encode(
        y=alt.Y("tipo_servico:N", sort="-x", title="Tipo de serviço"),
        x=alt.X("qtd:Q", title="Quantidade"),
        color=alt.Color("resultado:N", scale=color_scale),
        tooltip=["tipo_servico", "resultado", "qtd"]
    ).properties(height=360)

    st.altair_chart(chart_sla, use_container_width=True)
    st.dataframe(sla_t, use_container_width=True, height=240)

with tab4:
    st.subheader("Base detalhada (filtrada)")
    st.dataframe(f, use_container_width=True, height=520)
    st.download_button(
        "⬇️ Baixar base filtrada (CSV)",
        data=f.to_csv(index=False).encode("utf-8"),
        file_name="base_filtrada.csv",
        mime="text/csv"
    )
